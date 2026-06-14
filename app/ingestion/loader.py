"""
Document loading utilities — format-aware, extensible.

Architecture:
  Private _extract_* functions take raw bytes + source label and return Documents.
  Public load_*_file functions read from disk and delegate to the matching extractor.
  load_file(path) is the router: picks the right loader based on file extension.
  load_from_url(url) is the smart URL loader: probes Content-Type via HEAD first,
    falls back to URL extension, downloads, optionally saves binary docs, then parses.
  load_from_bytes(content, filename) handles uploaded file bytes the same way.

Supported formats:
  .pdf   — text extraction per page (pypdf)
  .docx  — paragraphs + table cells (python-docx)
  .md / .markdown / .txt / .rst — plain text passthrough
  .html / .htm  — BeautifulSoup with semantic-container preference
  .csv   — rows converted to "key: value" lines
  .json  — pretty-printed JSON object(s)
  .jsonl — newline-delimited JSON records
"""

import csv
import io
import json
import logging
import re
from pathlib import Path
from typing import Callable
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

_HEADERS = {"User-Agent": "RAG-Doc-Assistant/1.0 (research)"}


# ── Private format extractors ────────────────────────────────────────────────
# All accept (data: bytes | str, source: str) and return list[Document].
# Imports for optional heavy packages are deferred to the function body so
# the server still starts even if a package is missing (error surfaces only
# when that format is actually requested).


def _extract_pdf(data: bytes, source: str) -> list[Document]:
    """Extract text from a PDF, one Document per page."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("pypdf is required for PDF support: pip install pypdf")

    reader = PdfReader(io.BytesIO(data))
    docs: list[Document] = []
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": source, "page": page_num},
            ))
    logger.info("PDF %s → %d pages extracted", source, len(docs))
    return docs


def _extract_docx(data: bytes, source: str) -> list[Document]:
    """Extract paragraphs and table cells from a Word document."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required for DOCX support: pip install python-docx")

    doc = DocxDocument(io.BytesIO(data))

    parts: list[str] = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n\n".join(parts)
    if not text.strip():
        logger.warning("DOCX %s — no extractable text", source)
        return []

    logger.info("DOCX %s → %d chars", source, len(text))
    return [Document(page_content=text, metadata={"source": source})]


def _extract_html(data: bytes | str, source: str) -> list[Document]:
    """Parse HTML and return the main text content."""
    soup = BeautifulSoup(data, "html.parser")

    # Prefer semantic containers; fall back to full body
    container = (
        soup.find("main")
        or soup.find("article")
        or soup.find("div", {"class": re.compile(r"content|docs|documentation", re.I)})
        or soup.body
    )
    text = container.get_text(separator="\n", strip=True) if container else ""

    if not text.strip():
        logger.warning("HTML %s — no extractable text", source)
        return []

    logger.info("HTML %s → %d chars", source, len(text))
    return [Document(page_content=text, metadata={"source": source})]


def _extract_text(data: bytes | str, source: str) -> list[Document]:
    """Load plain text (Markdown, RST, or raw .txt)."""
    text = data.decode("utf-8", errors="replace") if isinstance(data, bytes) else data
    if not text.strip():
        logger.warning("Text %s — empty content", source)
        return []
    logger.info("Text %s → %d chars", source, len(text))
    return [Document(page_content=text, metadata={"source": source})]


def _extract_csv(data: bytes, source: str) -> list[Document]:
    """Convert CSV rows to 'key: value' lines, capped at 500 rows."""
    text_data = data.decode("utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(text_data))
    lines: list[str] = []
    for i, row in enumerate(reader):
        if i >= 500:
            logger.warning("CSV %s — truncated at 500 rows", source)
            break
        pairs = ", ".join(f"{k}: {v}" for k, v in row.items() if v and v.strip())
        if pairs:
            lines.append(pairs)

    if not lines:
        logger.warning("CSV %s — no rows extracted", source)
        return []

    logger.info("CSV %s → %d rows", source, len(lines))
    return [Document(page_content="\n".join(lines), metadata={"source": source})]


def _extract_json(data: bytes, source: str) -> list[Document]:
    """Pretty-print a JSON object or array as a Document."""
    try:
        obj = json.loads(data)
    except json.JSONDecodeError:
        logger.exception("Invalid JSON in %s", source)
        return []

    text = json.dumps(obj, indent=2, ensure_ascii=False)
    logger.info("JSON %s → %d chars", source, len(text))
    return [Document(page_content=text, metadata={"source": source})]


def _extract_jsonl(data: bytes, source: str) -> list[Document]:
    """Parse newline-delimited JSON (one object per line), capped at 200 records."""
    lines = data.decode("utf-8", errors="replace").splitlines()
    records: list[str] = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        try:
            records.append(json.dumps(json.loads(line), indent=2))
        except json.JSONDecodeError:
            pass
        if len(records) >= 200:
            logger.warning("JSONL %s — truncated at 200 records", source)
            break

    if not records:
        logger.warning("JSONL %s — no valid records", source)
        return []

    logger.info("JSONL %s → %d records", source, len(records))
    return [Document(page_content="\n---\n".join(records), metadata={"source": source})]


# ── Extension and MIME type routing maps ─────────────────────────────────────

_EXT_MAP: dict[str, Callable[[bytes, str], list[Document]]] = {
    ".pdf":      _extract_pdf,
    ".docx":     _extract_docx,
    ".doc":      _extract_docx,   # legacy Word — python-docx handles many .doc files
    ".html":     _extract_html,
    ".htm":      _extract_html,
    ".md":       _extract_text,
    ".markdown": _extract_text,
    ".txt":      _extract_text,
    ".rst":      _extract_text,
    ".csv":      _extract_csv,
    ".json":     _extract_json,
    ".jsonl":    _extract_jsonl,
    ".ndjson":   _extract_jsonl,
}

_MIME_MAP: dict[str, Callable[[bytes, str], list[Document]]] = {
    "application/pdf":               _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    "application/msword":            _extract_docx,
    "text/html":                     _extract_html,
    "application/xhtml+xml":         _extract_html,
    "text/plain":                    _extract_text,
    "text/markdown":                 _extract_text,
    "text/x-rst":                    _extract_text,
    "text/csv":                      _extract_csv,
    "application/csv":               _extract_csv,
    "application/json":              _extract_json,
    "application/x-ndjson":         _extract_jsonl,
    "application/jsonl":             _extract_jsonl,
}

# Formats whose raw bytes are worth persisting locally when fetched via URL
_BINARY_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
_BINARY_EXTENSIONS = {".pdf", ".docx", ".doc"}


# ── Public per-format file loaders ───────────────────────────────────────────

def load_pdf_file(path: Path) -> list[Document]:
    """Load a PDF file, returning one Document per page."""
    return _extract_pdf(path.read_bytes(), path.name)


def load_docx_file(path: Path) -> list[Document]:
    """Load a DOCX/DOC file, returning paragraphs and table cells."""
    return _extract_docx(path.read_bytes(), path.name)


def load_markdown_file(path: Path) -> list[Document]:
    """Load a Markdown file as a single Document (alias for load_text_file)."""
    return _extract_text(path.read_bytes(), path.name)


def load_text_file(path: Path) -> list[Document]:
    """Load a plain-text file (.txt, .rst, etc.) as a single Document."""
    return _extract_text(path.read_bytes(), path.name)


def load_html_file(path: Path) -> list[Document]:
    """Load a local HTML file, stripping tags and boilerplate."""
    return _extract_html(path.read_bytes(), path.name)


def load_csv_file(path: Path) -> list[Document]:
    """Load a CSV file, converting rows to readable key: value lines."""
    return _extract_csv(path.read_bytes(), path.name)


def load_json_file(path: Path) -> list[Document]:
    """Load a JSON file (.json or .jsonl) as a Document."""
    if path.suffix.lower() in (".jsonl", ".ndjson"):
        return _extract_jsonl(path.read_bytes(), path.name)
    return _extract_json(path.read_bytes(), path.name)


def load_file(path: Path) -> list[Document]:
    """
    Route to the correct format loader based on file extension.
    Falls back to plain-text for unknown extensions.
    """
    ext = path.suffix.lower()
    loader = _EXT_MAP.get(ext)
    if loader is None:
        logger.warning("Unknown extension '%s' for %s — treating as plain text", ext, path.name)
        loader = _extract_text
    return loader(path.read_bytes(), path.name)


# ── URL loader with content-type detection ────────────────────────────────────

def _url_filename(url: str) -> str:
    """Derive a safe local filename from a URL."""
    name = Path(urlparse(url).path).name
    if not name:
        name = urlparse(url).netloc.replace(".", "_") or "download"
    # Keep only safe characters
    name = re.sub(r"[^\w.\-]", "_", name)
    return name or "download"


def load_from_url(url: str, save_dir: Path | None = None, timeout: int = 30) -> list[Document]:
    """
    Smart URL loader:

    1. HEAD request to read Content-Type (avoids downloading large files blindly).
    2. Map MIME type → extractor; fall back to URL file extension; default to HTML.
    3. GET request to download the document.
    4. If the content is a binary document (PDF/DOCX) and save_dir is set, persist
       the raw bytes to disk before parsing.
    5. Parse with the detected extractor and return Documents.
    """
    # 1. Detect content type
    content_type = ""
    try:
        head = requests.head(url, headers=_HEADERS, timeout=timeout, allow_redirects=True)
        raw_ct = head.headers.get("Content-Type", "")
        content_type = raw_ct.split(";")[0].strip().lower()
    except Exception as exc:
        logger.debug("HEAD failed for %s: %s — will infer from URL extension", url, exc)

    # 2. Resolve extractor
    extractor = _MIME_MAP.get(content_type)
    if extractor is None:
        ext = Path(urlparse(url).path).suffix.lower()
        extractor = _EXT_MAP.get(ext, _extract_html)  # HTML is the web default

    is_binary = (
        content_type in _BINARY_MIME_TYPES
        or Path(urlparse(url).path).suffix.lower() in _BINARY_EXTENSIONS
    )

    # 3. Download
    try:
        resp = requests.get(url, headers=_HEADERS, timeout=timeout)
        resp.raise_for_status()
    except Exception:
        logger.exception("Failed to fetch %s", url)
        return []

    data = resp.content
    logger.info(
        "Downloaded %s — %d bytes, content-type: %s",
        url, len(data), content_type or "unknown",
    )

    # 4. Save binary docs to disk
    if is_binary and save_dir is not None:
        save_dir = Path(save_dir)
        save_dir.mkdir(parents=True, exist_ok=True)
        dest = save_dir / _url_filename(url)
        dest.write_bytes(data)
        logger.info("Saved binary doc → %s", dest)

    # 5. Parse
    return extractor(data, url)


# ── Upload handler ────────────────────────────────────────────────────────────

def load_from_bytes(content: bytes, filename: str) -> list[Document]:
    """
    Load an uploaded file from raw bytes.
    Routes to the appropriate extractor based on the filename extension.
    """
    ext = Path(filename).suffix.lower()
    extractor = _EXT_MAP.get(ext)
    if extractor is None:
        logger.warning("Unknown extension '%s' in upload '%s' — treating as plain text", ext, filename)
        extractor = _extract_text
    logger.info("Loading upload '%s' as %s (%d bytes)", filename, ext or "text", len(content))
    return extractor(content, filename)
